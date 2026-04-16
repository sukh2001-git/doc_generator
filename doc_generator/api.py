import frappe
from frappe import _
import os
from html.parser import HTMLParser


class TermsParser(HTMLParser):
    def __init__(self):
        super().__init__()
        self.result = []
        self._current_tag = None
        self._buf = ""
        self._in_ui = False

    def handle_starttag(self, tag, attrs):
        attr_dict = dict(attrs)
        if "ql-ui" in attr_dict.get("class", ""):
            self._in_ui = True
            return
        if tag in ("h1", "h2", "h3", "h4"):
            self._current_tag = "heading"
            self._buf = ""
        elif tag == "li":
            self._current_tag = "item"
            self._buf = ""

    def handle_endtag(self, tag):
        import html as htmllib
        if tag in ("h1", "h2", "h3", "h4") and self._current_tag == "heading":
            text = htmllib.unescape(self._buf.strip())
            if text:
                self.result.append({"type": "heading", "text": text})
            self._current_tag = None
            self._buf = ""
        elif tag == "li" and self._current_tag == "item":
            text = htmllib.unescape(self._buf.strip())
            if text:
                self.result.append({"type": "item", "text": text})
            self._current_tag = None
            self._buf = ""
        elif tag == "span":
            self._in_ui = False

    def handle_data(self, data):
        if not self._in_ui and self._current_tag in ("heading", "item"):
            self._buf += data


def _parse_terms(html_str):
    parser = TermsParser()
    parser.feed(html_str or "")
    return parser.result


def _get_settings():
    s = frappe.get_single("Doc Generator Settings")
    return {
        "fallback_logo": s.fallback_logo or None,
        "subtitle":      s.document_subtitle or "System Integrator Quotation",
        "brand_color":   s.brand_color or "#1A5276",
        "footer_text":   s.footer_text or "Generated from OneHash ERP",
    }


def _hex_to_rgb(hex_color: str):
    from docx.shared import RGBColor
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


@frappe.whitelist()
def generate_si_docx(quotation_name: str, offers: str = None) -> dict:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_ALIGN_VERTICAL
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        frappe.throw(_("python-docx is not installed. Run: bench pip install python-docx"))

    doc      = frappe.get_doc("Quotation", quotation_name)
    settings = _get_settings()

    BRAND_BLUE = _hex_to_rgb(settings["brand_color"])
    BRAND_HEX  = settings["brand_color"].lstrip("#")
    WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
    DARK_TEXT  = RGBColor(0x1A, 0x1A, 0x2E)
    MID_GREY   = RGBColor(0x55, 0x55, 0x55)

    document = Document()

    for section in document.sections:
        section.top_margin = section.bottom_margin = Inches(1)
        section.left_margin = section.right_margin = Inches(1)

    document.styles["Normal"].font.name = "Calibri"
    document.styles["Normal"].font.size = Pt(11)

    def set_cell_bg(cell, hex_color):
        tcPr = cell._tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color.upper())
        tcPr.append(shd)

    def add_hr(para, color=None, size=6):
        color = color or BRAND_HEX
        pPr  = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    str(size))
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), color)
        pBdr.append(bot)
        pPr.append(pBdr)

    def remove_cell_borders(cell):
        tcPr = cell._tc.get_or_add_tcPr()
        tcBdr = OxmlElement("w:tcBdr")
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            bdr = OxmlElement(f"w:{side}")
            bdr.set(qn("w:val"),   "none")
            bdr.set(qn("w:sz"),    "0")
            bdr.set(qn("w:space"), "0")
            bdr.set(qn("w:color"), "auto")
            tcBdr.append(bdr)
        tcPr.append(tcBdr)

    def section_title(text):
        p = document.add_paragraph(text)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(6)
        r = p.runs[0]
        r.bold           = True
        r.font.size      = Pt(11)
        r.font.color.rgb = BRAND_BLUE
        add_hr(p)

    # ── Header ───────────────────────────────────────────────────────────────
    import urllib.request, tempfile

    company_doc = frappe.get_doc("Company", doc.company)
    logo_url    = None
    if company_doc.company_logo:
        logo_url = company_doc.company_logo
    elif settings["fallback_logo"]:
        logo_url = settings["fallback_logo"]
    if logo_url and logo_url.startswith("/"):
        logo_url = frappe.utils.get_url(logo_url)

    header_tbl = document.add_table(rows=1, cols=2)
    for cell in header_tbl.rows[0].cells:
        remove_cell_borders(cell)

    logo_cell = header_tbl.rows[0].cells[0]
    logo_cell.width = Inches(2.0)
    logo_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    try:
        if logo_url:
            suffix = ".webp" if "webp" in logo_url else ".png"
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                urllib.request.urlretrieve(logo_url, tmp.name)
                lp = logo_cell.paragraphs[0]
                lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
                lp.add_run().add_picture(tmp.name, width=Inches(1.8))
    except Exception:
        # Logo failed to load, leave cell empty
        pass

    info_cell = header_tbl.rows[0].cells[1]
    info_cell.width = Inches(4.75)
    info_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = info_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(settings["subtitle"])
    r.font.size      = Pt(14)
    r.font.color.rgb = BRAND_BLUE

    hr_para = document.add_paragraph()
    hr_para.paragraph_format.space_before = Pt(6)
    hr_para.paragraph_format.space_after  = Pt(6)
    add_hr(hr_para, size=18)

    # Meta row
    meta = document.add_paragraph()
    meta.paragraph_format.space_before = Pt(6)
    meta.paragraph_format.space_after  = Pt(14)

    def meta_chunk(label, value):
        r = meta.add_run(f"{label}: ")
        r.bold           = True
        r.font.size      = Pt(10)
        r.font.color.rgb = BRAND_BLUE
        r2 = meta.add_run(f"{value}    ")
        r2.font.size      = Pt(10)
        r2.font.color.rgb = DARK_TEXT

    meta_chunk("Date", str(doc.transaction_date))
    if doc.valid_till:
        meta_chunk("Valid Till", str(doc.valid_till))

    # ── Customer Details ──────────────────────────────────────────────────────
    section_title("CUSTOMER DETAILS")
    
    # Simple customer name display
    cust_para = document.add_paragraph()
    cust_para.paragraph_format.space_before = Pt(4)
    cust_para.paragraph_format.space_after = Pt(12)
    cust_para.paragraph_format.left_indent = Inches(0.2)
    
    cust_run = cust_para.add_run(doc.customer_name or doc.party_name or "—")
    cust_run.bold = True
    cust_run.font.size = Pt(13)
    cust_run.font.color.rgb = DARK_TEXT

    document.add_paragraph()

    # ── Items Table(s) - Multiple Offers ─────────────────────────────────────
    import json
    
    # Parse offers if provided, otherwise create single offer from main items
    if offers:
        offers_data = json.loads(offers)
    else:
        # Fallback to main items table
        offers_data = [{
            'title': 'Items & Services',
            'items': doc.items,
            'net_total': doc.net_total,
            'taxes': doc.total_taxes_and_charges,
            'grand_total': doc.grand_total
        }]
    
    # Generate a table for each offer
    for offer_idx, offer in enumerate(offers_data):
        # Add spacing between offers
        if offer_idx > 0:
            document.add_paragraph()
            spacer = document.add_paragraph()
            spacer.paragraph_format.space_before = Pt(10)
            spacer.paragraph_format.space_after = Pt(10)
        
        section_title(offer.get('title', f'Offer {offer_idx + 1}'))
        
        # Get items from offer
        offer_items = offer.get('items', [])
        
        # Check if we need GST column for this offer
        has_gst = any(item.get('gst_amount') for item in offer_items)
        
        if has_gst:
            COLS   = ["#", "Item Code", "Item Name", "Qty", "UOM", "Rate", "GST", "Amount"]
            WIDTHS = [0.30, 1.00, 2.00, 0.50, 0.50, 0.90, 0.60, 1.00]
        else:
            COLS   = ["#", "Item Code", "Item Name", "Qty", "UOM", "Rate", "Amount"]
            WIDTHS = [0.35, 1.10, 2.20, 0.55, 0.55, 1.00, 1.00]

        itbl = document.add_table(rows=1, cols=len(COLS))
        itbl.style = "Table Grid"
        
        # Set column widths
        for i, w in enumerate(WIDTHS):
            itbl.columns[i].width = Inches(w)

        # Header row with proper alignment
        for i, name in enumerate(COLS):
            cell = itbl.rows[0].cells[i]
            set_cell_bg(cell, BRAND_HEX)
            r = cell.paragraphs[0].add_run(name)
            r.bold           = True
            r.font.size      = Pt(9)
            r.font.color.rgb = WHITE
            
            # Set alignment based on column
            if i == 0:  # # column - center
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif i >= 3:  # Qty, UOM, Rate, GST, Amount - right align
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:  # Item Code, Item Name - left align
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        for idx, item in enumerate(offer_items):
            row = itbl.add_row()
            if idx % 2 == 1:
                for cell in row.cells:
                    set_cell_bg(cell, "F9FBFC")

            def wc(ci, text, align=WD_ALIGN_PARAGRAPH.LEFT, bold=False):
                p = row.cells[ci].paragraphs[0]
                r = p.add_run(str(text or ""))
                r.font.size = Pt(9)
                if bold:
                    r.bold = True
                p.alignment = align

            wc(0, idx + 1, WD_ALIGN_PARAGRAPH.CENTER)
            wc(1, item.get('item_code', ''))
            wc(2, item.get('item_name', ''))
            wc(3, item.get('qty', 0), WD_ALIGN_PARAGRAPH.RIGHT)
            wc(4, item.get('uom', ''), WD_ALIGN_PARAGRAPH.LEFT)
            wc(5, frappe.utils.fmt_money(item.get('rate', 0), currency=doc.currency), WD_ALIGN_PARAGRAPH.RIGHT)
            
            if has_gst:
                gst_val = item.get('gst_amount', 0) or 0
                wc(6, frappe.utils.fmt_money(gst_val, currency=doc.currency), WD_ALIGN_PARAGRAPH.RIGHT)
                wc(7, frappe.utils.fmt_money(item.get('amount', 0), currency=doc.currency), WD_ALIGN_PARAGRAPH.RIGHT, bold=True)
            else:
                wc(6, frappe.utils.fmt_money(item.get('amount', 0), currency=doc.currency), WD_ALIGN_PARAGRAPH.RIGHT, bold=True)

        def total_row(label, value, bold=False, highlight=False):
            row = itbl.add_row()
            merged = row.cells[0].merge(row.cells[len(COLS) - 2])
            p = merged.paragraphs[0]
            r = p.add_run(label)
            r.bold      = bold
            r.font.size = Pt(10) if bold else Pt(9)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            val_cell = row.cells[len(COLS) - 1]
            if highlight:
                set_cell_bg(val_cell, "EAF2FB")
            
            vp  = val_cell.paragraphs[0]
            rv  = vp.add_run(str(value))
            rv.bold      = bold
            rv.font.size = Pt(10) if bold else Pt(9)
            vp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Get totals from offer
        net_total = offer.get('net_total', 0)
        taxes = offer.get('taxes', 0)
        grand_total = offer.get('grand_total', 0)
        
        total_row("Net Total", frappe.utils.fmt_money(net_total, currency=doc.currency))
        if taxes:
            total_row("Taxes", frappe.utils.fmt_money(taxes, currency=doc.currency))
        total_row("Grand Total", frappe.utils.fmt_money(grand_total, currency=doc.currency), bold=True, highlight=True)

    document.add_paragraph()

    # ── Terms & Conditions ────────────────────────────────────────────────────
    section_title("TERMS & CONDITIONS")
    terms_data = _parse_terms(doc.terms or "")

    if not terms_data:
        p = document.add_paragraph("No terms and conditions specified.")
        p.paragraph_format.left_indent = Inches(0.2)
        for r in p.runs:
            r.font.size = Pt(10)
    else:
        counter = 0
        for entry in terms_data:
            if entry["type"] == "heading":
                counter = 0
                p = document.add_paragraph(entry["text"])
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after  = Pt(3)
                p.paragraph_format.left_indent  = Inches(0.2)
                r = p.runs[0]
                r.bold           = True
                r.font.size      = Pt(10)
                r.font.color.rgb = BRAND_BLUE
            else:
                counter += 1
                p = document.add_paragraph()
                p.paragraph_format.left_indent  = Inches(0.4)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(3)
                p.add_run(f"{counter}. ").bold = True
                p.runs[0].font.size = Pt(10)
                txt = p.add_run(entry["text"])
                txt.font.size = Pt(10)

    # ── Footer ────────────────────────────────────────────────────────────────
    fp = document.add_paragraph()
    add_hr(fp, color="CCCCCC")
    fp.paragraph_format.space_before = Pt(10)
    r = fp.add_run(f"{settings['footer_text']}  ·  {doc.company}  ·  {frappe.utils.now()}")
    r.font.size      = Pt(8)
    r.font.color.rgb = MID_GREY
    fp.alignment     = WD_ALIGN_PARAGRAPH.CENTER

    # ── Save ──────────────────────────────────────────────────────────────────
    filename  = f"SI_Doc_{frappe.scrub(quotation_name)}.docx"
    file_path = os.path.join(frappe.get_site_path("public", "files"), filename)
    os.makedirs(os.path.dirname(file_path), exist_ok=True)
    document.save(file_path)

    return {"file_url": f"/files/{filename}"}